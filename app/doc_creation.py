def generate_docx_from_segments(transcription_file, template_file, output_file):
    import re
    from docx import Document
    import requests

    gemini_api_key = "AIzaSyBVzPQY8m7wg1iweqE2ZMXHyF0I_0NaXxs"
    endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={gemini_api_key}"

    with open(transcription_file, "r", encoding="utf-8") as f:
        content = f.read()

    pattern = re.compile(
    r"ODJ\s+(\d+)\s+[-—]\s+(.+?)\r?\nRésumé\s*:\s*(.*?)(?=\r?\nODJ\s+\d+\s+[-—]|\Z)",
    re.DOTALL
    )

    segments_raw = pattern.findall(content)
    print(f"DEBUG: segments_raw = {segments_raw}")

    segments = []
    for num, title, text in segments_raw:
        segments.append({
            "number": int(num),
            "title": title.strip(),
            "text": text.strip()
        })

    print(f"🔍 Segments trouvés : {len(segments)}")
    print(f"DEBUG: segments = {segments}")

    summaries, names = {}, {}
    headers = {"Content-Type": "application/json"}

    for seg in segments:
        num = seg["number"]
        raw = seg["text"]
        prompt = (
            f"Tu es secrétaire de réunion chez Junior ISEP, tu dois constituer un résumé clair de ce qui s’est dis et décidé pendant le Conseil d’Administration ; "
            "Être rédigé de façon neutre, impersonnelle et respectueuse des avis de chacun. "
            "Tu dois limiter le titre des ODJ à 2 à 5 mots maximum"
            "Tu dois être concis et garder l'essentiel des informations importantes et des décisions prises.\n\n"
            f"Texte source :\n\"\"\"\n{raw}\n\"\"\"\n\n"
            "Réponds au format :\nTitre : <titre résumé>\nRésumé : <texte résumé>"
        )

        payload = {"contents": [{"parts": [{"text": prompt}]}]}
        try:
            resp = requests.post(endpoint, headers=headers, json=payload)
            resp.raise_for_status()
            data = resp.json()
            output = data["candidates"][0]["content"]["parts"][0]["text"].strip()
            match = re.search(r"Titre\s*:\s*(.*?)\nRésumé\s*:\s*(.*)", output, flags=re.DOTALL)
            name_clean = match.group(1).strip() if match else seg["title"]
            summary_clean = match.group(2).strip() if match else "[Résumé manquant]"
        except Exception as e:
            print(f"❌ Erreur pour ODJ {num} : {e}")
            name_clean = seg["title"]
            summary_clean = "[Résumé manquant]"

        names[num] = name_clean
        summaries[num] = summary_clean

    def replace_placeholder_in_paragraph(paragraph, placeholder, replacement):
        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder not in full_text:
            return
        new_text = full_text.replace(placeholder, replacement)
        for run in paragraph.runs:
            run.text = ''
        paragraph.runs[0].text = new_text

    doc = Document(template_file)

    mapping = {f"Nom de l'ODJ {num}": names[num] for num in names}
    mapping.update({f"Contenu résumé de l'ODJ {num}": summaries[num] for num in summaries})


    for para in doc.paragraphs:
        for ph, val in mapping.items():
            token = f"{{{ph}}}"
            replace_placeholder_in_paragraph(para, token, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for ph, val in mapping.items():
                        token = f"{{{ph}}}"
                        replace_placeholder_in_paragraph(para, token, val)

    doc.save(output_file)
    print(f"📁 Document généré : {output_file}")
